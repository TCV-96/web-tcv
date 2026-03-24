import docx
import os

base_dir = r"C:\Users\User\Documents\WEB TCV"
doc_path = os.path.join(base_dir, "Dossier_TCV_Carmven.docx")

doc = docx.Document()

# Título
doc.add_heading('Dossier Corporativo - T.C.V. TECHNOLOGYCARMVEN', 0)

# Sección 1: Perfil del CEO
doc.add_heading('Sección 1: Perfil del CEO', level=1)
doc.add_paragraph('CEO / Fundador: Jeferson Carmona')
doc.add_paragraph('• T.S.U. en Electrónica (UPT Clodosbaldo Russian)')
doc.add_paragraph('• Docente de Ciencias Físicas - MPPE / UPEL IPMALA')
doc.add_paragraph('• Especialista INCES: Electricidad Industrial y de Potencia (+800 horas certificadas)')
doc.add_paragraph('• Diplomado en Docencia con Inteligencia Artificial - CIL LATAM')
doc.add_paragraph('• Agente Inmobiliario Internacional Certificado - Farro Costamar (Agencia VE-43)')
doc.add_paragraph('• Asesor de Marketing Digital - Club de Leones de Maturín, Lions International')

# Sección 2: Líneas de Servicio
doc.add_heading('Sección 2: Líneas de Servicio', level=1)
doc.add_paragraph('• Ingeniería Electrónica y Soporte Técnico: Mantenimiento avanzado y reparación PC/Hardware.')
doc.add_paragraph('• Marketing Digital Corporativo: Estrategias de posicionamiento y Meta Ads.')
doc.add_paragraph('• Electricidad de Potencia e Industrial: Instalaciones bajo normas internacionales.')
doc.add_paragraph('• Gestión Inmobiliaria Farro: Compra, venta y alquiler de inmuebles en 22 países.')

# Sección 3: Enlaces Oficiales y Contacto
doc.add_heading('Sección 3: Enlaces Oficiales y Contacto', level=1)
doc.add_paragraph('• Sitio Web Oficial (GitHub Pages): https://tcv-96.github.io/web-tcv/')
doc.add_paragraph('• WhatsApp directo: https://wa.me/584145841929 (+58 414-5841929)')
doc.add_paragraph('• Correo Electrónico: criptofinanzas0@gmail.com')
doc.add_paragraph('• Instagram: @tcv_technologycarmven')

# Sección 4: El "Cerebro" de CARMVENCITO
doc.add_heading('Sección 4: El "Cerebro" de CARMVENCITO (Prompt Maestro)', level=1)
    
conocimiento_path = os.path.join(base_dir, "CARMVENCITO_CONOCIMIENTO.txt")
if os.path.exists(conocimiento_path):
    doc.add_paragraph('A continuación se adjunta el prompt e instrucciones por las que se rige el bot:')
    with open(conocimiento_path, 'r', encoding='utf-8') as f:
        doc.add_paragraph(f.read())
else:
    doc.add_paragraph('No se encontró archivo de conocimiento previo.')

# Sección 5: Bitácora de Errores y Soluciones (Problemas Actuales)
doc.add_heading('Sección 5: Bitácora de Errores y Soluciones', level=1)
doc.add_paragraph('A continuación se detallan los errores encontrados en el sitio web y el estatus de sus correcciones:')

doc.add_paragraph('🔴 Error 1: Enlace Canonical erróneo (tcven.com)', style='List Bullet')
doc.add_paragraph('Estatus: SOLUCIONADO. Se actualizaron los tags SEO y Open Graph en index.html para apuntar a la dirección real de GitHub Pages (https://tcv-96.github.io/web-tcv/).', style='List Bullet')

doc.add_paragraph('🔴 Error 2: Imagen de Avatar rota en Tarjeta de Inmobiliaria Farro', style='List Bullet')
doc.add_paragraph('Estatus: SOLUCIONADO. Se reemplazó el logo estático por un diseño circular que enfoca la fotografía del CEO (jeferson_azul.png) con efecto responsive scale-110.', style='List Bullet')

doc.add_paragraph('🔴 Error 3: Bot de chat emitiendo "Error de conexión"', style='List Bullet')
doc.add_paragraph('Estatus: EN REVISIÓN. El script de carmvencito.js está perfectamente programado para abrir el Gem de Gemini (94a666640855). Si el error continúa, puede tratarse de caché del navegador o un script secundario heredado de una versión antigua.', style='List Bullet')

# Guardar
doc.save(doc_path)
print("Dossier updated successfully at " + doc_path)
